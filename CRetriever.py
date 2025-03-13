import os
import re
import json
import time
import xmlrpc.client
import docx
import tempfile
from docx.shared import Pt
from io import BytesIO

import openai
# OpenAI API key will be loaded from secrets

from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import streamlit as st
from bs4 import BeautifulSoup, NavigableString, Tag

# ------------------------------------------------------------------
# Get all settings from Streamlit secrets
# ------------------------------------------------------------------
def get_settings():
    """
    Get all application settings from Streamlit secrets with fallbacks to environment variables.
    Returns a dictionary with all the required configuration.
    """
    settings = {
        # Odoo settings
        "odoo_url": "",
        "odoo_db": "",
        "odoo_username": "",
        "odoo_password": "",
        
        # Google redirect URI
        "redirect_uri": "",
        
        # OpenAI API key
        "openai_api_key": "",
        
        # OpenAI model
        "openai_model": "gpt-4-turbo"
    }
    
    # Get Odoo settings
    if "odoo" in st.secrets:
        settings["odoo_url"] = st.secrets["odoo"].get("ODOO_URL", "")
        settings["odoo_db"] = st.secrets["odoo"].get("ODOO_DB", "")
        settings["odoo_username"] = st.secrets["odoo"].get("ODOO_USERNAME", "")
        settings["odoo_password"] = st.secrets["odoo"].get("ODOO_PASSWORD", "")
    
    # Get OAuth redirect URI
    if "auth" in st.secrets:
        settings["redirect_uri"] = st.secrets["auth"].get("redirect_uri", "")
    
    # Get OpenAI settings
    if "openai" in st.secrets:
        settings["openai_api_key"] = st.secrets["openai"].get("api_key", "")
        settings["openai_model"] = st.secrets["openai"].get("model", "gpt-4-turbo")
    
    # Fall back to environment variables if not found in secrets
    if not settings["odoo_url"]:
        settings["odoo_url"] = os.getenv("ODOO_URL", "")
    if not settings["odoo_db"]:
        settings["odoo_db"] = os.getenv("ODOO_DB", "")  
    if not settings["odoo_username"]:
        settings["odoo_username"] = os.getenv("ODOO_USERNAME", "")
    if not settings["odoo_password"]:
        settings["odoo_password"] = os.getenv("ODOO_PASSWORD", "")
    if not settings["redirect_uri"]:
        settings["redirect_uri"] = os.getenv("REDIRECT_URI", "")
    if not settings["openai_api_key"]:
        settings["openai_api_key"] = os.getenv("OPENAI_API_KEY", "")
    
    return settings


# ------------------------------------------------------------------
# Initialize OpenAI client
# ------------------------------------------------------------------
def setup_openai():
    settings = get_settings()
    openai.api_key = settings["openai_api_key"]
    return settings["openai_model"]


# ------------------------------------------------------------------
# Briefing Function using OpenAI
# ------------------------------------------------------------------
def summarize_text(text):
    # Set up OpenAI and get the model name
    model = setup_openai()
    
    if not openai.api_key:
        st.error("OpenAI API key not configured. Please add it to your secrets.toml file.")
        return "Unable to generate summary: API key missing"
    
    prompt = (
        "Please provide a comprehensive briefing of the following text. "
        "The briefing should particularly emphasize the detailed descriptions from the Odoo tasks and "
        "the sales order order lines. Include all key details and structure your briefing with the following headings:\n\n"
        "**Project**: Provide an overview of the project, including customer and order date.\n"
        "**Odoo Task Descriptions**: Focus on the task details, deadlines, and descriptions provided in Odoo.\n"
        "**Sales Order Order Lines**: Include detailed information about each order line (product, description, quantity, delivered).\n"
        "**Gmail Communications**: Summarize the key points from the Gmail data.\n"
        "**Conclusion**: Provide final insights.\n\n"
        "Please retain any hyperlinks as they appear. Here is the text for the briefing:\n\n" + text
    )
    
    try:
        response = openai.ChatCompletion.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=700
        )
        return response.choices[0].message["content"].strip()
    except Exception as e:
        st.error(f"Error generating summary: {str(e)}")
        return f"Error generating summary: {str(e)}"


# ------------------------------------------------------------------
# Advanced HTML Parsing to Markdown-like Text
# ------------------------------------------------------------------
def parse_html_to_text(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    
    def parse_node(node, indent=0, list_type=None):
        lines = []
        if isinstance(node, NavigableString):
            text = node.strip()
            if text:
                lines.append(" " * indent + text)
            return lines
        if isinstance(node, Tag):
            if node.name in ("ul", "ol"):
                for child in node.children:
                    lines.extend(parse_node(child, indent, node.name))
            elif node.name == "li":
                bullet = "- " if list_type != "ol" else "1. "
                sublines = []
                for child in node.children:
                    sublines.extend(parse_node(child, indent + 2, None))
                if sublines:
                    first_line = bullet + sublines[0].lstrip()
                    lines.append(" " * indent + first_line)
                    for line in sublines[1:]:
                        lines.append(" " * (indent + 2) + line)
                else:
                    lines.append(" " * indent + bullet)
            elif node.name == "a":
                link_text = node.get_text(strip=True)
                link_href = node.get("href", "")
                md_link = f"[{link_text}]({link_href})" if link_href else link_text
                lines.append(" " * indent + md_link)
            elif node.name == "br":
                lines.append("")
            elif node.name in ("p", "div"):
                sublines = []
                for child in node.children:
                    sublines.extend(parse_node(child, indent, list_type))
                if sublines:
                    lines.extend(sublines)
                    lines.append("")
            else:
                for child in node.children:
                    lines.extend(parse_node(child, indent, list_type))
        return lines

    parsed_lines = []
    for element in soup.children:
        parsed_lines.extend(parse_node(element, 0, None))
    final_lines = [line.rstrip() for line in parsed_lines if line.strip() != ""]
    return "\n".join(final_lines)


# ------------------------------------------------------------------
# Fetch Gmail Data (with OAuth authentication flow)
# ------------------------------------------------------------------
def get_gmail_service():
    """
    Handles Gmail authentication flow through Google OAuth.
    Uses Streamlit session state to maintain credentials between reruns.
    """
    # Get settings
    settings = get_settings()
    redirect_uri = settings["redirect_uri"]
    
    if not redirect_uri:
        st.error("Redirect URI not configured. Please add it to your secrets.toml file.")
        st.stop()
    
    # Debug tab to help troubleshoot authentication issues
    with st.expander("Authentication Debugging (Expand if having issues)"):
        st.write("Session state keys:", list(st.session_state.keys()))
        query_params = st.query_params
        st.write("Query parameters:", query_params)
        st.write(f"Using redirect URI: {redirect_uri}")
    
    # Check if we already have credentials
    if "gmail_creds" in st.session_state:
        creds = st.session_state.gmail_creds
        # Refresh token if expired
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
                st.session_state.gmail_creds = creds
            except Exception as e:
                st.error(f"Error refreshing credentials: {e}")
                # Clear credentials to restart auth flow
                del st.session_state.gmail_creds
                st.rerun()
    else:
        # Load client config from Streamlit secrets
        try:
            client_config_str = st.secrets["gcp"]["client_config"]
            client_config = json.loads(client_config_str)
            
            # Write the client config to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".json") as temp:
                temp.write(json.dumps(client_config).encode("utf-8"))
                temp_path = temp.name
            
            # Create the OAuth flow with the redirect URI from settings
            flow = InstalledAppFlow.from_client_secrets_file(
                temp_path,
                ['https://www.googleapis.com/auth/gmail.readonly'],
                redirect_uri=redirect_uri
            )
            
            # Check for authorization code in query parameters
            query_params = st.query_params
            if "code" in query_params:
                try:
                    # Get the authorization code
                    code = query_params["code"]
                    st.write("Attempting to exchange code for token...")
                    
                    # Exchange code for tokens
                    flow.fetch_token(code=code)
                    st.session_state.gmail_creds = flow.credentials
                    
                    # Clean up the URL by removing the query parameters
                    # Note: This might not work in all Streamlit environments
                    try:
                        st.experimental_set_query_params()
                    except:
                        pass
                        
                    st.success("Authentication successful!")
                    time.sleep(1)  # Give a moment for the success message to display
                    st.rerun()  # Rerun to clear the auth parameters from URL
                except Exception as e:
                    st.error(f"Error exchanging code for token: {str(e)}")
                    st.write("Please try again.")
                    # Generate a new authorization URL
                    auth_url, _ = flow.authorization_url(prompt='consent')
                    st.markdown(f"[Click here to authenticate with Google]({auth_url})")
                    st.stop()
            else:
                # No code parameter, start the auth flow
                auth_url, _ = flow.authorization_url(prompt='consent', access_type='offline')
                st.warning("You need to authenticate with Google to access your emails.")
                st.markdown(f"[Click here to authenticate with Google]({auth_url})")
                st.stop()
        except Exception as e:
            st.error(f"Error during authentication setup: {str(e)}")
            st.write("Please check your configuration and try again.")
            st.stop()

    try:
        # Build the Gmail service with our credentials
        service = build('gmail', 'v1', credentials=st.session_state.gmail_creds)
        return service
    except Exception as e:
        st.error(f"Error building Gmail service: {str(e)}")
        # Clear credentials to restart auth flow
        if "gmail_creds" in st.session_state:
            del st.session_state.gmail_creds
        st.stop()


def get_gmail_data(sales_order_code):
    """Fetch Gmail messages related to the sales order code"""
    service = get_gmail_service()
    
    query = f"{sales_order_code}"
    try:
        results = service.users().messages().list(userId='me', q=query).execute()
        formatted_messages = []
        
        if 'messages' in results:
            for i, msg in enumerate(results['messages']):
                # Show progress for large result sets
                if i % 5 == 0:
                    st.write(f"Processing email {i+1}/{len(results['messages'])}...")
                    
                message = service.users().messages().get(userId='me', id=msg['id']).execute()
                headers = message.get('payload', {}).get('headers', [])
                sender = "Unknown"
                subject = "No Subject"
                
                for header in headers:
                    if header.get("name", "").lower() == "from":
                        sender = header.get("value", "Unknown")
                    elif header.get("name", "").lower() == "subject":
                        subject = header.get("value", "No Subject")
                
                snippet = message.get('snippet', '')
                formatted = f"Sender: {sender}\nSubject: {subject}\nEmail content: {snippet}\n-----------------\n"
                formatted_messages.append(formatted)
            
            return "\n".join(formatted_messages)
        else:
            return "No emails found related to this sales order."
    except Exception as e:
        st.error(f"Error fetching Gmail data: {str(e)}")
        return f"Error fetching Gmail data: {str(e)}"


# ------------------------------------------------------------------
# Cached Odoo Connection
# ------------------------------------------------------------------
@st.cache_resource(show_spinner=False)
def get_odoo_connection(odoo_url, db_name, username, password):
    if not odoo_url or not db_name or not username or not password:
        raise Exception("Odoo connection details missing. Please check your configuration.")
        
    try:
        common = xmlrpc.client.ServerProxy(f'{odoo_url}/xmlrpc/2/common')
        uid = common.authenticate(db_name, username, password, {})
        if not uid:
            raise Exception("Failed to authenticate to Odoo. Check your credentials.")
        models = xmlrpc.client.ServerProxy(f'{odoo_url}/xmlrpc/2/object')
        return uid, models
    except Exception as e:
        raise Exception(f"Odoo connection failed: {str(e)}")


# ------------------------------------------------------------------
# Fetch Odoo Project Tasks (Filtered by Order Reference and Company)
# ------------------------------------------------------------------
def get_odoo_data(order_reference, company):
    settings = get_settings()
    odoo_url = settings["odoo_url"]
    odoo_db = settings["odoo_db"]
    odoo_username = settings["odoo_username"]
    odoo_password = settings["odoo_password"]
    
    try:
        uid, models = get_odoo_connection(odoo_url, odoo_db, odoo_username, odoo_password)
        domain = [
            ('sale_line_id.order_id.name', '=', order_reference),
            ('sale_line_id.company_id.name', '=', company)
        ]
        fields = ['id', 'name', 'user_ids', 'date_deadline', 'sale_line_id', 'description']
        all_tasks = models.execute_kw(
            odoo_db, uid, odoo_password, 'project.task', 'search_read', [domain],
            {'fields': fields}
        )
        def get_user_names(user_ids):
            if not user_ids:
                return []
            user_records = models.execute_kw(
                odoo_db, uid, odoo_password, 'res.users', 'search_read',
                [[('id', 'in', user_ids)]],
                {'fields': ['name']}
            )
            return [record['name'] for record in user_records]
        task_data = []
        seen_normalized = set()
        for task in all_tasks:
            user_ids = task.get('user_ids', [])
            user_names = get_user_names(user_ids)
            assigned_to = ', '.join(user_names) if user_names else 'N/A'
            sale_line = task.get('sale_line_id', [])
            sale_line_val = sale_line[1] if sale_line and isinstance(sale_line, list) and len(sale_line) == 2 else 'N/A'
            raw_desc = task.get('description', '') or ''
            desc_clean = parse_html_to_text(raw_desc)
            desc_text = f"Description:\n{desc_clean}\n" if desc_clean.strip() else ""
            task_info = (
                f"Task: {task.get('name', '')}\n"
                f"Assigned to: {assigned_to}\n"
                f"Deadline: {task.get('date_deadline', 'No deadline')}\n"
                f"Sale Line: {sale_line_val}\n"
                f"{desc_text}"
                "-----------------------------------------\n"
            )
            normalized = " ".join(task_info.strip().split())
            if normalized not in seen_normalized:
                seen_normalized.add(normalized)
                task_data.append(task_info)
        return "\n".join(task_data)
    except Exception as e:
        st.error(f"Error fetching Odoo tasks: {str(e)}")
        return f"Error fetching Odoo tasks: {str(e)}"


# ------------------------------------------------------------------
# Fetch Sales Order Details (with Order Lines) as Markdown
# ------------------------------------------------------------------
def get_sale_order_details(order_reference):
    settings = get_settings()
    odoo_url = settings["odoo_url"]
    odoo_db = settings["odoo_db"]
    odoo_username = settings["odoo_username"]
    odoo_password = settings["odoo_password"]
    
    try:
        uid, models = get_odoo_connection(odoo_url, odoo_db, odoo_username, odoo_password)
        orders = models.execute_kw(
            odoo_db, uid, odoo_password, 'sale.order', 'search_read',
            [[('name', '=', order_reference)]],
            {'fields': ['id', 'name', 'partner_id', 'date_order']}
        )
        if not orders:
            return "No sales order found with that reference."
        order = orders[0]
        sale_order_id = order.get('id')
        md_details = f"**Sales Order:** {order.get('name', '')}\n\n"
        if order.get('partner_id') and isinstance(order.get('partner_id'), list):
            md_details += f"**Customer:** {order['partner_id'][1]}\n\n"
        md_details += f"**Order Date:** {order.get('date_order', '')}\n\n"
        md_details += "### Order Lines:\n"
        order_lines = models.execute_kw(
            odoo_db, uid, odoo_password, 'sale.order.line', 'search_read',
            [[('order_id', '=', sale_order_id)]],
            {'fields': ['product_id', 'name', 'product_uom_qty', 'qty_delivered']}
        )
        for line in order_lines:
            product = line.get('product_id')[1] if line.get('product_id') and isinstance(line.get('product_id'), list) else "N/A"
            description = line.get('name', '')
            quantity = line.get('product_uom_qty', '')
            delivered = line.get('qty_delivered', '')
            md_details += f"- **Product:** {product}\n"
            md_details += f"  - **Description:** {description}\n"
            md_details += f"  - **Quantity:** {quantity}\n"
            md_details += f"  - **Delivered:** {delivered}\n\n"
        return md_details
    except Exception as e:
        st.error(f"Error fetching sales order details: {str(e)}")
        return f"Error fetching sales order details: {str(e)}"


# ------------------------------------------------------------------
# Fetch Order References for Dropdown (Filtered by Company)
# ------------------------------------------------------------------
def get_odoo_order_references(company):
    settings = get_settings()
    odoo_url = settings["odoo_url"]
    odoo_db = settings["odoo_db"]
    odoo_username = settings["odoo_username"]
    odoo_password = settings["odoo_password"]
    
    try:
        uid, models = get_odoo_connection(odoo_url, odoo_db, odoo_username, odoo_password)
        domain = [('company_id.name', '=', company)]
        lines = models.execute_kw(
            odoo_db, uid, odoo_password, 'sale.order.line', 'search_read', [domain],
            {'fields': ['order_id']}
        )
        order_refs = set()
        for line in lines:
            order_field = line.get('order_id')
            if order_field and isinstance(order_field, list) and len(order_field) == 2:
                order_refs.add(order_field[1])
        return sorted(list(order_refs))
    except Exception as e:
        st.error(f"Error fetching order references: {str(e)}")
        return []


# ------------------------------------------------------------------
# Fetch Companies from Sale Order Lines
# ------------------------------------------------------------------
def get_odoo_companies():
    settings = get_settings()
    odoo_url = settings["odoo_url"]
    odoo_db = settings["odoo_db"]
    odoo_username = settings["odoo_username"]
    odoo_password = settings["odoo_password"]
    
    try:
        uid, models = get_odoo_connection(odoo_url, odoo_db, odoo_username, odoo_password)
        lines = models.execute_kw(
            odoo_db, uid, odoo_password, 'sale.order.line', 'search_read', [[]],
            {'fields': ['company_id']}
        )
        companies = set()
        for line in lines:
            comp_field = line.get('company_id')
            if comp_field and isinstance(comp_field, list) and len(comp_field) == 2:
                companies.add(comp_field[1])
        return sorted(list(companies))
    except Exception as e:
        st.error(f"Error fetching companies: {str(e)}")
        return []


# ------------------------------------------------------------------
# Create Word Document In-Memory with Regex-Based Bold Conversion
# ------------------------------------------------------------------
def create_word_document_in_memory(doc_title, content_text):
    try:
        bold_pattern = re.compile(r"\*\*(.+?)\*\*")
        doc_io = BytesIO()
        doc = docx.Document()
        doc.add_heading(doc_title, 0)
        lines = content_text.split("\n")
        for line in lines:
            paragraph = doc.add_paragraph()
            last_end = 0
            for match in bold_pattern.finditer(line):
                start, end = match.span()
                normal_text = line[last_end:start]
                if normal_text:
                    run_normal = paragraph.add_run(normal_text)
                    run_normal.bold = False
                bold_text = match.group(1)
                run_bold = paragraph.add_run(bold_text)
                run_bold.bold = True
                last_end = end
            if last_end < len(line):
                trailing_text = line[last_end:]
                run_trailing = paragraph.add_run(trailing_text)
                run_trailing.bold = False
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        # Return an empty BytesIO object as fallback
        empty_io = BytesIO()
        empty_io.seek(0)
        return empty_io.getvalue()


# ------------------------------------------------------------------
# Streamlit App
# ------------------------------------------------------------------
def main():
    st.set_page_config(
        page_title="Sales Order Data Retriever",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    st.title("📊 Sales Order Data Retriever")
    
    # Add sidebar with info and logout button
    with st.sidebar:
        st.header("About")
        st.write(
            "This app retrieves and summarizes data from multiple sources "
            "including Odoo, Gmail, and generates an AI-powered summary."
        )
        
        # Add logout option
        if "gmail_creds" in st.session_state:
            if st.button("Logout from Google"):
                del st.session_state.gmail_creds
                st.success("Logged out successfully!")
                st.rerun()
    
    # Get settings to check Odoo connection
    settings = get_settings()
    
    # Check that Odoo settings are properly configured
    if not settings["odoo_url"] or not settings["odoo_db"] or not settings["odoo_username"] or not settings["odoo_password"]:
        st.error("Odoo connection details missing. Please check your configuration.")
        st.info("Add Odoo details to your secrets.toml file or as environment variables.")
        
        # Show a sample secrets.toml configuration
        with st.expander("Sample secrets.toml configuration"):
            st.code("""
[odoo]
ODOO_URL = "https://your-odoo-instance.com/"
ODOO_DB = "your-database-name"
ODOO_USERNAME = "your-username"
ODOO_PASSWORD = "your-password"

[gcp]
client_config = '{"web":{"client_id":"your-client-id.apps.googleusercontent.com","project_id":"your-project-id","auth_uri":"https://accounts.google.com/o/oauth2/auth","token_uri":"https://oauth2.googleapis.com/token","auth_provider_x509_cert_url":"https://www.googleapis.com/oauth2/v1/certs","client_secret":"your-client-secret"}}'

[openai]
api_key = "your-openai-api-key"
model = "gpt-4-turbo"

[auth]
redirect_uri = "https://your-streamlit-app-url/"
            """, language="toml")
        
        st.stop()

    # Step 1: Select Company
    with st.spinner("Fetching Companies from Odoo..."):
        try:
            companies = get_odoo_companies()
        except Exception as e:
            st.error(f"Error fetching companies: {e}")
            companies = []
    
    if companies:
        placeholder_company = "-- Select a Company --"
        company_options = [placeholder_company] + companies
        selected_company = st.selectbox("Select a Company", company_options, index=0)
        if selected_company == placeholder_company:
            st.info("Please select a valid company.")
            st.stop()
    else:
        st.warning("No companies found.")
        st.info("Make sure your Odoo connection is correctly configured.")
        selected_company = None
        st.stop()

    # Step 2: Select Order Reference (filtered by company)
    with st.spinner("Fetching Order References from Odoo..."):
        try:
            order_refs = get_odoo_order_references(selected_company)
        except Exception as e:
            st.error(f"Error fetching order references: {e}")
            order_refs = []
    
    if order_refs:
        placeholder_order = "-- Select an Order Reference --"
        order_options = [placeholder_order] + order_refs
        selected_order_ref = st.selectbox("Select an Order Reference", order_options, index=0)
        if selected_order_ref == placeholder_order:
            st.info("Please select a valid order reference.")
            st.stop()
    else:
        st.warning("No order references found for the selected company.")
        selected_order_ref = None
        st.stop()

    # Step 3: Fetch Data
    if st.button("Fetch Data", type="primary") and selected_order_ref and selected_company:
        # Create placeholders for each section
        gmail_placeholder = st.empty()
        odoo_placeholder = st.empty()
        sales_placeholder = st.empty()
        summary_placeholder = st.empty()
        
        # Initialize data
        gmail_data = ""
        odoo_tasks = ""
        sale_order_details = ""
        
        gmail_placeholder.info("Fetching Gmail data...")
        odoo_placeholder.info("Waiting to fetch Odoo tasks...")
        sales_placeholder.info("Waiting to fetch Sales Order details...")
        summary_placeholder.info("Waiting to generate summary...")
        
        # Fetch Gmail data
        try:
            gmail_data = get_gmail_data(selected_order_ref)
            gmail_placeholder.success("Gmail data fetched successfully!")
        except Exception as e:
            gmail_placeholder.error(f"Error fetching Gmail data: {str(e)}")
            gmail_data = f"Error fetching Gmail data: {str(e)}"
        
        # Fetch Odoo tasks
        odoo_placeholder.info("Fetching Odoo tasks...")
        try:
            odoo_tasks = get_odoo_data(selected_order_ref, selected_company)
            odoo_placeholder.success("Odoo tasks fetched successfully!")
        except Exception as e:
            odoo_placeholder.error(f"Error fetching Odoo tasks: {str(e)}")
            odoo_tasks = f"Error fetching Odoo tasks: {str(e)}"
        
        # Fetch Sales Order details
        sales_placeholder.info("Fetching Sales Order details...")
        try:
            sale_order_details = get_sale_order_details(selected_order_ref)
            sales_placeholder.success("Sales Order details fetched successfully!")
        except Exception as e:
            sales_placeholder.error(f"Error fetching Sales Order details: {str(e)}")
            sale_order_details = f"Error fetching Sales Order details: {str(e)}"
        
        # Combine all data for summary
        combined_text = (
            f"Gmail Data:\n{gmail_data}\n\n"
            f"Odoo Tasks:\n{odoo_tasks}\n\n"
            f"Sales Order Details:\n{sale_order_details}"
        )
        
        # Generate AI summary
        summary_placeholder.info("Generating AI briefing...")
        try:
            summary = summarize_text(combined_text)
            st.session_state.summary_text = summary
            summary_placeholder.success("AI briefing generated successfully!")
        except Exception as e:
            summary_placeholder.error(f"Error generating summary: {str(e)}")
            summary = f"Error generating summary: {str(e)}"
            st.session_state.summary_text = summary

        # Clear the placeholders
        gmail_placeholder.empty()
        odoo_placeholder.empty()
        sales_placeholder.empty()
        summary_placeholder.empty()

        # Reorder tabs: Summary first, then Sales Order Details, then Odoo Tasks, then Gmail Data
        tabs = st.tabs(["Summary", "Sales Order Details", "Odoo Tasks", "Gmail Data"])

        with tabs[0]:
            st.header("Summary")
            st.markdown(st.session_state.summary_text)
            st.download_button(
                label="Download Summary as Word",
                data=create_word_document_in_memory("Summary", st.session_state.summary_text),
                file_name="Summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with tabs[1]:
            st.header("Sales Order Details")
            st.markdown(sale_order_details)
            st.download_button(
                label="Download Sales Order Details as Word",
                data=create_word_document_in_memory("Sales Order Details", sale_order_details),
                file_name="Sales_Order_Details.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with tabs[2]:
            st.header("Odoo Tasks")
            if odoo_tasks.strip():
                for line in odoo_tasks.split("\n"):
                    st.write(line)
            else:
                st.write("No Odoo tasks found.")
            st.download_button(
                label="Download Odoo Tasks as Word",
                data=create_word_document_in_memory("Odoo Tasks", odoo_tasks),
                file_name="Odoo_Tasks.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with tabs[3]:
            st.header("Gmail Data")
            if gmail_data.strip():
                for line in gmail_data.split("\n"):
                    st.write(line)
            else:
                st.write("No Gmail data found.")
            st.download_button(
                label="Download Gmail Data as Word",
                data=create_word_document_in_memory("Gmail Data", gmail_data),
                file_name="Gmail_Data.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.success("All data retrieved and processed successfully!")


if __name__ == "__main__":
    main()